import subprocess, math2docx, re

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pypdf import PdfWriter, PdfReader

from umk import Umk
from translation import translate

def generateDokladnoiFromWord(name, date, faculty, dean, department, group, lesson_name, lesson_type, start_time, end_time):
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''
    hdr_cells[2].text = f'Декану {faculty}\n{dean}'
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = document.add_paragraph(f'\n\n\n\n\n\n\n\n\n\nДокладная\n')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph(f'Довожу до Вашего сведения, что группа {group} не явилась на {lesson_type} занятия с {start_time} по {end_time} {date} числа по дисциплине "{lesson_name}".')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph('\n\n\n\n')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = f'Преподаватель кафедры {department}'
    hdr_cells[1].text = ''
    hdr_cells[2].text = f'{name}'
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save('output/word/dokladnoi.docx')

    generate_pdf("output/word/dokladnoi.docx", 'output/word')

def generateUmkFromWord(token):
    umkDoc = Umk.query.filter_by(token=token).first()
    lang = umkDoc.language

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    p = document.add_paragraph(translate('ministry_name', lang) + '\n\n'
                                + translate('university_name', lang) + '\n\n'
                                + umkDoc.faculty + '\n\n'
                                + umkDoc.department + '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n')

    if lang == 'kz':
        p.add_run('"' + umkDoc.subject + '"\n')
        p.add_run(translate('discipline', lang) + '\n\n')
        p.add_run(umkDoc.group + '\n')
        p.add_run(translate('ed_program', lang) + '\n\n')
        p.add_run(translate('document_name', lang) + '\n\n\n\n\n\n\n\n\n\n').bold = True
    else:
        p.add_run(translate('document_name', lang) + '\n\n').bold = True
        p.add_run(translate('discipline', lang) + '\n')
        p.add_run('"' + umkDoc.subject + '"\n\n')
        p.add_run(translate('ed_program', lang) + '\n')
        p.add_run(umkDoc.group + '\n\n\n\n\n\n\n\n\n\n')

    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = 1

    p = document.add_paragraph(translate('course', lang) + ': ' + umkDoc.course + '\n'
                               + translate('duration', lang) + ': ' + umkDoc.studyTime + '\n'
                               + translate('credits', lang) + ': ' + umkDoc.credits + '\n\n\n\n\n\n\n\n\n\n'
                               + translate('department_approval_part_one', lang) + '\n'
                               + translate('department_approval_part_two', lang).format(department_protocol = umkDoc.department_protocol, department_date = umkDoc.department_date) + '\n\n'
                               + translate('faculty_meeting_part_one', lang) + '\n'
                               + translate('faculty_meeting_part_two', lang).format(faculty_protocol = umkDoc.faculty_protocol, faculty_date = umkDoc.faculty_date))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    document.add_page_break()

    p = document.add_paragraph(translate('thematic_plan', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    themes = umkDoc.first_themes.split('/;-;/ ')
    lections = umkDoc.first_lections.split('/;-;/ ')
    seminars = umkDoc.first_seminars.split('/;-;/ ')
    labs = umkDoc.first_labs.split('/;-;/ ')
    srsps = umkDoc.first_srsps.split('/;-;/ ')
    srss = umkDoc.first_srss.split('/;-;/ ')

    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(60)
    table.columns[2].width = Mm(20)
    table.columns[3].width = Mm(20)
    table.columns[4].width = Mm(20)
    table.columns[5].width = Mm(20)
    table.columns[6].width = Mm(20)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('t_name', lang)
    hdr_cells[2].text = translate('t_lec', lang)
    hdr_cells[3].text = translate('t_prac', lang)
    hdr_cells[4].text = translate('t_lab', lang)
    hdr_cells[5].text = translate('t_iwst', lang)
    hdr_cells[6].text = translate('t_iws', lang)
    for i in range(umkDoc.first_counts):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = themes[i]
        row_cells[2].text = lections[i]
        row_cells[3].text = seminars[i]
        row_cells[4].text = labs[i]
        row_cells[5].text = srsps[i]
        row_cells[6].text = srss[i]

    row_cells = table.add_row().cells
    row_cells[0].text = ''
    row_cells[1].text = translate('t_total', lang)
    row_cells[2].text = str(sum(map(int, list(filter(None, lections)))))
    row_cells[3].text = str(sum(map(int, list(filter(None, seminars)))))
    row_cells[4].text = str(sum(map(int, list(filter(None, labs)))))
    row_cells[5].text = str(sum(map(int, list(filter(None, srsps)))))
    row_cells[6].text = str(sum(map(int, list(filter(None, srss)))))

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('teacher_data', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    lecturers = umkDoc.second_lecturers.split('/;-;/ ')

    for i in lecturers:
        p = document.add_paragraph(i)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(22.7)
        p.paragraph_format.line_spacing = 1
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('course_policy', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('course_policy_text', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('course_prerequisites', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(umkDoc.second_prerequisites)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('course_postrequisites', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(umkDoc.second_post_requisites)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('description', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    for i in umkDoc.third_course_description.split('\n'):
        p = document.add_paragraph(i)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(22.7)
        p.paragraph_format.line_spacing = 1
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('outcomes_and_methods', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    loss = umkDoc.third_los.split('/;-;/ ')
    methods = umkDoc.third_methods.split('/;-;/ ')

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(60)
    table.columns[2].width = Mm(100)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('lo_outcomes', lang)
    hdr_cells[2].text = translate('lo_methods', lang)
    for i in range(umkDoc.third_counts):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = loss[i]
        row_cells[2].text = methods[i]

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('teaching_methods', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(umkDoc.fourth_teaching_methods)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('methods_for_lo', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(umkDoc.fourth_grade_methods)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('sources', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    bls = umkDoc.fifth_bls.split('/;-;/ ')
    als = umkDoc.fifth_als.split('/;-;/ ')

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(160)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('sources_course', lang)

    row_cells = table.add_row().cells
    mergedCell = row_cells[1].merge(row_cells[0])
    mergedCell.text = translate('sources_basic', lang)
    mergedCell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(umkDoc.fifth_bl_counts):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = bls[i]

    row_cells = table.add_row().cells
    mergedCell = row_cells[1].merge(row_cells[0])
    mergedCell.text = translate('sources_additional', lang)
    mergedCell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(umkDoc.fifth_al_counts):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = als[i]

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('lecture', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    lss = umkDoc.sixth_lss.split('/;-;/ ')
    lps = umkDoc.sixth_lps.split('/;-;/ ')

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(60)
    table.columns[2].width = Mm(100)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('lecture_topic', lang)
    hdr_cells[2].text = translate('lecture_plan', lang)
    for i in range(umkDoc.sixth_counts):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = lss[i]
        row_cells[2].text = lps[i]

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('practical', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    spss = umkDoc.seventh_spss.split('/;-;/ ')
    spqs = umkDoc.seventh_spqs.split('/;-;/ ')
    sprs = umkDoc.seventh_sprs.split('/;-;/ ')
    spls = umkDoc.seventh_spls.split('/;-;/ ')

    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(30)
    table.columns[2].width = Mm(70)
    table.columns[3].width = Mm(40)
    table.columns[4].width = Mm(20)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('practical_topic', lang)
    hdr_cells[2].text = translate('practical_question', lang)
    hdr_cells[3].text = translate('practical_recommendation', lang)
    hdr_cells[4].text = translate('practical_link', lang)
    for i in range(umkDoc.seventh_counts):
        spqsArray = re.split('\[:(.*?)\:]', spqs[i])
        spqsModifiedArray = re.sub('\[:(.*?)\:]', '-:-formula-:-', spqs[i]).split('-:-')
        isSpqsFirstLine = True
        sprsArray = re.split('\[:(.*?)\:]', sprs[i])
        sprsModifiedArray = re.sub('\[:(.*?)\:]', '-:-formula-:-', sprs[i]).split('-:-')
        isSprsFirstLine = True
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = spss[i]
        for j in range(len(spqsArray)):
            if spqsModifiedArray[j] == '' or spqsModifiedArray[j] == '\n': continue
            if isSpqsFirstLine:
                if spqsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[2].paragraphs[0], spqsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[2].text = spqsArray[j].strip('\n')
                isSpqsFirstLine = False
            else:
                if spqsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[2].add_paragraph(), spqsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[2].add_paragraph(spqsArray[j].strip('\n'))
        for j in range(len(sprsArray)):
            if sprsModifiedArray[j] == '' or sprsModifiedArray[j] == '\n': continue
            if isSprsFirstLine:
                if sprsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[3].paragraphs[0], sprsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[3].text = sprsArray[j].strip('\n')
                isSprsFirstLine = False
            else:
                if sprsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[3].add_paragraph(), sprsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[3].add_paragraph(sprsArray[j].strip('\n'))
        row_cells[4].text = spls[i]

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('laboratory', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    if umkDoc.eighth_counts is not None and umkDoc.eighth_counts > 0:
        slabs = umkDoc.eighth_slabs.split('/;-;/ ')
        qlabs = umkDoc.eighth_qlabs.split('/;-;/ ')
        rlabs = umkDoc.eighth_rlabs.split('/;-;/ ')

        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Mm(10)
        table.columns[1].width = Mm(40)
        table.columns[2].width = Mm(80)
        table.columns[3].width = Mm(40)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = translate('laboratory_topic', lang)
        hdr_cells[2].text = translate('laboratory_task', lang)
        hdr_cells[3].text = translate('laboratory_recommendation', lang)
        for i in range(umkDoc.eighth_counts):
            qlabsArray = re.split('\[:(.*?)\:]', qlabs[i])
            qlabsModifiedArray = re.sub('\[:(.*?)\:]', '-:-formula-:-', qlabs[i]).split('-:-')
            isQlabsFirstLine = True
            rlabsArray = re.split('\[:(.*?)\:]', rlabs[i])
            rlabsModifiedArray = re.sub('\[:(.*?)\:]', '-:-formula-:-', rlabs[i]).split('-:-')
            isRlabsFirstLine = True
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = slabs[i]
            for j in range(len(qlabsArray)):
                if qlabsModifiedArray[j] == '' or qlabsModifiedArray[j] == '\n': continue
                if isQlabsFirstLine:
                    if qlabsModifiedArray[j] == 'formula':
                        math2docx.add_math(row_cells[2].paragraphs[0], qlabsArray[j].replace('\\\\', '\\'))
                    else:
                        row_cells[2].text = qlabsArray[j].strip('\n')
                    isQlabsFirstLine = False
                else:
                    if qlabsModifiedArray[j] == 'formula':
                        math2docx.add_math(row_cells[2].add_paragraph(), qlabsArray[j].replace('\\\\', '\\'))
                    else:
                        row_cells[2].add_paragraph(qlabsArray[j].strip('\n'))
            for j in range(len(rlabsArray)):
                if rlabsModifiedArray[j] == '' or rlabsModifiedArray[j] == '\n': continue
                if isRlabsFirstLine:
                    if rlabsModifiedArray[j] == 'formula':
                        math2docx.add_math(row_cells[3].paragraphs[0], rlabsArray[j].replace('\\\\', '\\'))
                    else:
                        row_cells[3].text = rlabsArray[j].strip('\n')
                    isRlabsFirstLine = False
                else:
                    if rlabsModifiedArray[j] == 'formula':
                        math2docx.add_math(row_cells[3].add_paragraph(), rlabsArray[j].replace('\\\\', '\\'))
                    else:
                        row_cells[3].add_paragraph(rlabsArray[j].strip('\n'))
    else:
        p = document.add_paragraph(translate('laboratory_not_provided', lang))
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(22.7)
        p.paragraph_format.line_spacing = 1
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('iwst', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    sropss = umkDoc.ninth_sropss.split('/;-;/ ')
    sropqs = umkDoc.ninth_sropqs.split('/;-;/ ')
    sroprs = umkDoc.ninth_sroprs.split('/;-;/ ')

    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(30)
    table.columns[2].width = Mm(80)
    table.columns[3].width = Mm(50)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('iwst_topic', lang)
    hdr_cells[2].text = translate('iwst_task', lang)
    hdr_cells[3].text = translate('iwst_recommendation', lang)
    for i in range(umkDoc.ninth_srop_counts):
        sropqsArray = re.split('\[:(.*?)\:]', sropqs[i])
        sropqsModifiedArray = re.sub('\[:(.*?)\:]', '-:-formula-:-', sropqs[i]).split('-:-')
        isSropqsFirstLine = True
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = sropss[i]
        for j in range(len(sropqsArray)):
            if sropqsModifiedArray[j] == '' or sropqsModifiedArray[j] == '\n': continue
            if isSropqsFirstLine:
                if sropqsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[2].paragraphs[0], sropqsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[2].text = sropqsArray[j].strip('\n')
                isSropqsFirstLine = False
            else:
                if sropqsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[2].add_paragraph(), sropqsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[2].add_paragraph(sropqsArray[j].strip('\n'))
        row_cells[3].text = sroprs[i]

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('iws', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    sross = umkDoc.ninth_sross.split('/;-;/ ')
    sroqs = umkDoc.ninth_sroqs.split('/;-;/ ')
    srors = umkDoc.ninth_srors.split('/;-;/ ')

    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(30)
    table.columns[2].width = Mm(90)
    table.columns[3].width = Mm(40)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('iws_topic', lang)
    hdr_cells[2].text = translate('iws_task', lang)
    hdr_cells[3].text = translate('iws_recommendation', lang)
    for i in range(umkDoc.ninth_sro_counts):
        sroqsArray = re.split('\[:(.*?)\:]', sroqs[i])
        sroqsModifiedArray = re.sub('\[:(.*?)\:]', '-:-formula-:-', sroqs[i]).split('-:-')
        isSroqsFirstLine = True
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = sross[i]
        for j in range(len(sroqsArray)):
            if sroqsModifiedArray[j] == '' or sroqsModifiedArray[j] == '\n': continue
            if isSroqsFirstLine:
                if sroqsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[2].paragraphs[0], sroqsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[2].text = sroqsArray[j].strip('\n')
                isSroqsFirstLine = False
            else:
                if sroqsModifiedArray[j] == 'formula':
                    math2docx.add_math(row_cells[2].add_paragraph(), sroqsArray[j].replace('\\\\', '\\'))
                else:
                    row_cells[2].add_paragraph(sroqsArray[j].strip('\n'))
        row_cells[3].text = srors[i]

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('paper_topics', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    pws = umkDoc.tenth_pws.split('/;-;/ ')

    for i in range(umkDoc.tenth_counts):
        p = document.add_paragraph(str(i + 1) + '. ' + pws[i])
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(22.7)
        p.paragraph_format.line_spacing = 1
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('evaluation_policy', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('evaluation_policy_text', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    gps = umkDoc.eleventh_gps.split('/;-;/ ')
    gpt = umkDoc.eleventh_gpt.split('/;-;/ ')
    gpf = umkDoc.eleventh_gpf.split('/;-;/ ')
    gpr = umkDoc.eleventh_gpr.split('/;-;/ ')
    gpd = umkDoc.eleventh_gpd.split('/;-;/ ')
    gpb = umkDoc.eleventh_gpb.split('/;-;/ ')

    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(35)
    table.columns[2].width = Mm(25)
    table.columns[3].width = Mm(25)
    table.columns[4].width = Mm(25)
    table.columns[5].width = Mm(25)
    table.columns[6].width = Mm(25)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('ep_theme', lang)
    hdr_cells[2].text = translate('ep_type_of_lesson', lang)
    hdr_cells[3].text = translate('ep_type_of_task', lang)
    hdr_cells[4].text = translate('ep_report', lang)
    hdr_cells[5].text = translate('ep_due_date', lang)
    hdr_cells[6].text = translate('ep_scores', lang)
    for i in range(umkDoc.eleventh_counts):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = gps[i]
        row_cells[2].text = gpt[i]
        row_cells[3].text = gpf[i]
        row_cells[4].text = gpr[i]
        row_cells[5].text = gpd[i]
        row_cells[6].text = gpb[i]

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    p = document.add_paragraph(translate('assessment_criteria', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    acLetters = ['A', 'A-', 'B+', 'B', 'B-', 'C+', 'C', 'C-', 'D+', 'D', 'Fx', 'F']
    acNumbers = ['4,0', '3,67', '3,33', '3,0', '2,67', '2,33', '2,0', '1,67', '1,33', '1,0', '0,5', '0']
    acProcents = ['95-100', '90-94', '85-89', '80-84', '75-79', '70-74', '65-69', '60-64', '55-59', '50-54', '25-49', '0-24']
    acTexts = umkDoc.twelfth_texts.split('/;-;/ ')

    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(25)
    table.columns[1].width = Mm(25)
    table.columns[2].width = Mm(25)
    table.columns[3].width = Mm(95)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = translate('ac_marking', lang)
    hdr_cells[1].text = translate('ac_points', lang)
    hdr_cells[2].text = translate('ac_percentage', lang)
    hdr_cells[3].text = translate('ac_criteria', lang)
    for i in range(len(acLetters)):
        row_cells = table.add_row().cells
        row_cells[0].text = acLetters[i]
        row_cells[1].text = acNumbers[i]
        row_cells[2].text = acProcents[i]
        row_cells[3].text = acTexts[i]

    document.save('output/word/umk.docx')

    generate_pdf("output/word/umk.docx", 'output/word')

    document.add_page_break()

    p = document.add_paragraph(translate('content', lang))
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(22.7)
    p.paragraph_format.line_spacing = 1

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(15)
    table.columns[1].width = Mm(140)
    table.columns[2].width = Mm(15)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = translate('content_name', lang)
    hdr_cells[2].text = translate('content_page', lang)
    i = 1
    for name, page in getHeadingsPageNumbers(lang).items():
        split = name.split(' ', 1)
        row_cells = table.add_row().cells
        row_cells[0].text = split[0]
        row_cells[1].text = split[1]
        row_cells[2].text = str(page)
        i = i + 1

    document.save('output/word/umk.docx')

    generate_pdf("output/word/umk.docx", 'output/word')

    # original = PdfReader('output/word/umk.pdf')
    # demo = original.pages[0]
    # sign = PdfReader('word/sign.pdf').pages[0]
    # demo.merge_page(sign)
    #
    # writer = PdfWriter()
    #
    # for page in original.pages:
    #     writer.add_page(page)
    #
    # with open("output/word/umk.pdf", "wb") as outFile:
    #     writer.write(outFile)

    return token

def getHeadingsPageNumbers(lang):
    resPages = {}

    for page in PdfReader('output/word/umk.pdf').pages:
        formattedText = page.extract_text().replace('\n', ' ').replace('  ', ' ')

        if formattedText.find(translate('thematic_plan', lang)) > -1: resPages[translate('thematic_plan', lang)] = page.page_number
        if formattedText.find(translate('teacher_data', lang)) > -1: resPages[translate('teacher_data', lang)] = page.page_number
        if formattedText.find(translate('course_policy', lang)) > -1: resPages[translate('course_policy', lang)] = page.page_number
        if formattedText.find(translate('course_prerequisites', lang)) > -1: resPages[translate('course_prerequisites', lang)] = page.page_number
        if formattedText.find(translate('course_postrequisites', lang)) > -1: resPages[translate('course_postrequisites', lang)] = page.page_number
        if formattedText.find(translate('description', lang)) > -1: resPages[translate('description', lang)] = page.page_number
        if formattedText.find(translate('outcomes_and_methods', lang)) > -1: resPages[translate('outcomes_and_methods', lang)] = page.page_number
        if formattedText.find(translate('teaching_methods', lang)) > -1: resPages[translate('teaching_methods', lang)] = page.page_number
        if formattedText.find(translate('methods_for_lo', lang)) > -1: resPages[translate('methods_for_lo', lang)] = page.page_number
        if formattedText.find(translate('sources', lang)) > -1: resPages[translate('sources', lang)] = page.page_number
        if formattedText.find(translate('lecture', lang)) > -1: resPages[translate('lecture', lang)] = page.page_number
        if formattedText.find(translate('practical', lang)) > -1: resPages[translate('practical', lang)] = page.page_number
        if formattedText.find(translate('laboratory', lang)) > -1: resPages[translate('laboratory', lang)] = page.page_number
        if formattedText.find(translate('iwst', lang)) > -1: resPages[translate('iwst', lang)] = page.page_number
        if formattedText.find(translate('iws', lang)) > -1: resPages[translate('iws', lang)] = page.page_number
        if formattedText.find(translate('paper_topics', lang)) > -1: resPages[translate('paper_topics', lang)] = page.page_number
        if formattedText.find(translate('evaluation_policy', lang)) > -1: resPages[translate('evaluation_policy', lang)] = page.page_number
        if formattedText.find(translate('assessment_criteria', lang)) > -1: resPages[translate('assessment_criteria', lang)] = page.page_number

    return resPages

def generate_pdf(doc_path, path):
    subprocess.call(['soffice',
                     # '--headless',
                     '--convert-to',
                     'pdf',
                     '--outdir',
                     path,
                     doc_path])

    return doc_path